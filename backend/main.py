import asyncio
import uuid
from typing import List, Dict, Any
from fastapi import FastAPI, HTTPException, Header, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from ai_service import generate_world_lore, chat_with_world, generate_image_from_prompt, generate_story_stream
from models import WorldLore
import database
import tempfile
import os
from docx import Document
from fpdf import FPDF
from starlette.background import BackgroundTask

app = FastAPI(title="Atlas Studio API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class SparkInput(BaseModel):
    spark: str
    genre: str | None = None

class RenameInput(BaseModel):
    name: str

class ChatInput(BaseModel):
    message: str

class ImageInput(BaseModel):
    prompt: str

@app.get("/")
def read_root():
    return {"message": "Welcome to the Atlas Studio Persistent API!"}

@app.get("/api/projects")
def get_projects(x_user_email: str | None = Header(default=None)):
    if not x_user_email:
        raise HTTPException(status_code=401, detail="Email required")
    projects = database.get_all_projects(x_user_email)
    return {"projects": projects}

@app.get("/api/project/{project_id}")
def get_project(project_id: str, x_user_email: str | None = Header(default=None)):
    proj = database.get_project(project_id)
    if not proj or proj.get("userEmail") != x_user_email:
        raise HTTPException(status_code=404, detail="Project not found")
    return proj

@app.post("/api/project/{project_id}/archive")
def archive_project(project_id: str, x_user_email: str | None = Header(default=None)):
    proj = database.get_project(project_id)
    if proj and proj.get("userEmail") == x_user_email:
        proj["isArchived"] = True
        database.save_project(proj)
    return {"success": True}

@app.delete("/api/project/{project_id}")
def delete_project(project_id: str, x_user_email: str | None = Header(default=None)):
    proj = database.get_project(project_id)
    if proj and proj.get("userEmail") == x_user_email:
        database.delete_project(project_id)
    return {"success": True}

@app.post("/api/project/{project_id}/rename")
def rename_project(project_id: str, input_data: RenameInput, x_user_email: str | None = Header(default=None)):
    proj = database.get_project(project_id)
    if proj and proj.get("userEmail") == x_user_email:
        proj["customName"] = input_data.name
        database.save_project(proj)
    return {"success": True}

@app.get("/api/project/{project_id}/download/docx")
def download_docx(project_id: str, x_user_email: str | None = Header(default=None)):
    proj = database.get_project(project_id)
    if not proj or proj.get("userEmail") != x_user_email:
        raise HTTPException(status_code=404, detail="Project not found")
        
    doc = Document()
    world_name = proj.get("customName") or (proj.get("lore", {}).get("world_name") if proj.get("lore") else "Untitled World")
    doc.add_heading(world_name, 0)
    
    if proj.get("lore"):
        lore = proj["lore"]
        doc.add_heading("World Context", level=1)
        doc.add_heading("Core Setting", level=2)
        doc.add_paragraph(lore.get("core_history", ""))
        doc.add_heading("Magic & Technology", level=2)
        doc.add_paragraph(lore.get("magic_system", ""))
        
    doc.add_heading("The Story", level=1)
    doc.add_paragraph(proj.get("storyContent", ""))
    
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    return FileResponse(path, filename=f"{world_name}.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", background=BackgroundTask(os.remove, path))

@app.get("/api/project/{project_id}/download/pdf")
def download_pdf(project_id: str, x_user_email: str | None = Header(default=None)):
    proj = database.get_project(project_id)
    if not proj or proj.get("userEmail") != x_user_email:
        raise HTTPException(status_code=404, detail="Project not found")
        
    world_name = proj.get("customName") or (proj.get("lore", {}).get("world_name") if proj.get("lore") else "Untitled World")
    
    pdf = FPDF()
    pdf.add_page()
    
    # Use standard Arial/Helvetica instead of custom font to avoid missing font issues
    pdf.set_font("Helvetica", size=16)
    
    # We use a simple helper to replace utf-8 chars not supported by latin-1 to avoid errors
    def sanitize(text):
        if not text: return ""
        return text.encode('latin-1', 'replace').decode('latin-1')

    pdf.cell(200, 10, txt=sanitize(world_name), ln=1, align="C")
    
    pdf.set_font("Helvetica", size=12)
    if proj.get("lore"):
        lore = proj["lore"]
        pdf.ln(10)
        pdf.set_font("Helvetica", style="B", size=14)
        pdf.cell(200, 10, txt="World Context", ln=1)
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 10, txt=sanitize(lore.get("core_history", "")))
        pdf.ln(5)
        pdf.multi_cell(0, 10, txt=sanitize(lore.get("magic_system", "")))
        
    pdf.ln(10)
    pdf.set_font("Helvetica", style="B", size=14)
    pdf.cell(200, 10, txt="The Story", ln=1)
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 10, txt=sanitize(proj.get("storyContent", "")))
    
    fd, path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    pdf.output(path)
    return FileResponse(path, filename=f"{world_name}.pdf", media_type="application/pdf", background=BackgroundTask(os.remove, path))

def run_generation_pipeline(project_id: str):
    proj = database.get_project(project_id)
    if not proj: return
    try:
        lore_obj = generate_world_lore(proj["spark"], proj.get("genre"))
        proj["lore"] = lore_obj.model_dump()
        database.save_project(proj)
        
        full_story = ""
        stream = generate_story_stream(lore_obj, proj.get("genre"))
        for i, chunk in enumerate(stream):
            proj = database.get_project(project_id)
            if proj.get("status") == "abort":
                proj["status"] = "stopped"
                database.save_project(proj)
                return
            
            full_story += chunk
            # Save to DB periodically to simulate streaming to polling clients
            if i % 3 == 0:
                proj["storyContent"] = full_story
                database.save_project(proj)
                
        # Final save
        proj = database.get_project(project_id)
        if proj.get("status") != "abort":
            proj["storyContent"] = full_story
            proj["status"] = "done"
            database.save_project(proj)
    except Exception as e:
        proj = database.get_project(project_id)
        proj["status"] = "error"
        proj["error"] = str(e)
        database.save_project(proj)

@app.post("/api/project/start")
def start_world(input_data: SparkInput, background_tasks: BackgroundTasks, x_user_email: str | None = Header(default=None)):
    if not x_user_email:
        raise HTTPException(status_code=401, detail="Email required")
    project_id = str(uuid.uuid4())
    project = {
        "id": project_id,
        "userEmail": x_user_email,
        "spark": input_data.spark,
        "genre": input_data.genre,
        "lore": None,
        "chatHistory": [{"role": "model", "content": "Welcome to Atlas Studio. I am the Story Weaver. What world shall we weave today?"}],
        "storyContent": "",
        "status": "generating",
        "isArchived": False
    }
    database.save_project(project)
    
    background_tasks.add_task(run_generation_pipeline, project_id)
    
    return {"project_id": project_id}

@app.post("/api/project/{project_id}/generate_lore")
def generate_lore_endpoint(project_id: str, x_user_email: str | None = Header(default=None)):
    proj = database.get_project(project_id)
    if not proj or proj.get("userEmail") != x_user_email:
        raise HTTPException(status_code=404, detail="Project not found")
        
    try:
        lore_obj = generate_world_lore(proj["spark"], proj.get("genre"))
        proj["lore"] = lore_obj.model_dump()
        database.save_project(proj)
        return proj["lore"]
    except Exception as e:
        proj["status"] = "error"
        proj["error"] = str(e)
        database.save_project(proj)
        raise HTTPException(status_code=500, detail=str(e))

def story_streamer(project_id: str, lore_obj: WorldLore):
    full_story = ""
    try:
        stream = generate_story_stream(lore_obj)
        for chunk in stream:
            proj = database.get_project(project_id)
            if proj.get("status") == "abort":
                proj["status"] = "stopped"
                database.save_project(proj)
                return # Stop generator
            
            full_story += chunk
            yield chunk
            
        # Final save
        proj = database.get_project(project_id)
        if proj.get("status") != "abort":
            proj["storyContent"] = full_story
            proj["status"] = "done"
            database.save_project(proj)
            
    except Exception as e:
        proj = database.get_project(project_id)
        proj["status"] = "error"
        proj["error"] = str(e)
        database.save_project(proj)

@app.post("/api/project/{project_id}/generate_story")
def generate_story_endpoint(project_id: str, x_user_email: str | None = Header(default=None)):
    proj = database.get_project(project_id)
    if not proj or proj.get("userEmail") != x_user_email:
        raise HTTPException(status_code=404, detail="Project not found")
        
    if not proj.get("lore"):
        raise HTTPException(status_code=400, detail="Lore not generated yet")
        
    lore_obj = WorldLore.model_validate(proj["lore"])
    return StreamingResponse(story_streamer(project_id, lore_obj), media_type="text/plain")

@app.post("/api/project/{project_id}/stop")
def stop_world(project_id: str, x_user_email: str | None = Header(default=None)):
    proj = database.get_project(project_id)
    if proj and proj.get("userEmail") == x_user_email and proj["status"] == "generating":
        proj["status"] = "abort"
        database.save_project(proj)
    return {"success": True}

@app.post("/api/project/{project_id}/chat")
def chat_project(project_id: str, input_data: ChatInput, x_user_email: str | None = Header(default=None)):
    proj = database.get_project(project_id)
    if not proj or proj.get("userEmail") != x_user_email:
        raise HTTPException(status_code=404, detail="Project not found")
        
    # Append user message
    proj["chatHistory"].append({"role": "user", "content": input_data.message})
    database.save_project(proj)
    
    # Convert dict back to WorldLore object
    lore_obj = None
    if proj.get("lore"):
        lore_obj = WorldLore.model_validate(proj["lore"])
        
    try:
        story_content = proj.get("storyContent", "")
        reply = chat_with_world(input_data.message, lore_obj, proj["chatHistory"], story_content)
        proj = database.get_project(project_id) # Refresh
        proj["chatHistory"].append({"role": "model", "content": reply})
        database.save_project(proj)
        return {"reply": reply}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/project/{project_id}/chat_image")
def chat_image_project(project_id: str, input_data: ImageInput, x_user_email: str | None = Header(default=None)):
    proj = database.get_project(project_id)
    if not proj or proj.get("userEmail") != x_user_email:
        raise HTTPException(status_code=404, detail="Project not found")
        
    proj["chatHistory"].append({"role": "user", "content": f"/imagine {input_data.prompt}"})
    database.save_project(proj)
    
    try:
        image_b64 = generate_image_from_prompt(input_data.prompt)
        proj = database.get_project(project_id)
        proj["chatHistory"].append({
            "role": "model", 
            "content": f'Generated image for: "{input_data.prompt}"',
            "image_url": image_b64
        })
        database.save_project(proj)
        return {"image_url": image_b64}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
